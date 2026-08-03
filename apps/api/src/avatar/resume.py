"""
Turn an uploaded resume into text an interviewer can be briefed from.

**Why the text is extracted once, at upload, and stored.** The alternative is parsing the file
on every session start, which puts a PDF parse inside the latency budget for no benefit and
makes the briefing depend on a library version rather than on a decision. Extracting once means
the text is inspectable in the console before anyone is interviewed against it — and a resume
that extracted badly is something an operator should see rather than discover from a strange
first question.

**What this deliberately does not do.** No parsing into fields. No "years of experience", no
skill list, no employer timeline. Every resume parser that tries produces confident nonsense on
the resumes that do not fit its assumptions, and a confidently wrong structured field is worse
than a paragraph of text — because the interviewer would state it as fact. The interviewer gets
the text and is told to treat it as the candidate's own claims, unverified.

**Truncation is stated, not silent.** A long resume is cut at `MAX_BRIEFING_CHARS` and the
record says how much was kept. Silently sending half a document to a model is the kind of thing
that explains a bad interview three weeks later.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

MAX_BRIEFING_CHARS = 12_000
"""
How much resume text reaches the interviewer.

Roughly three to four dense pages. Chosen against what a resume actually is rather than against
a context window: past this, a document is a portfolio or a scanned book, and the extra tokens
buy worse questions rather than better ones because the model's attention is spread over
material the candidate will not be asked about.
"""

SUPPORTED = (".pdf", ".docx", ".txt", ".md", ".markdown")

MAX_RESUME_BYTES = 10 * 1024 * 1024
"""
10 MB. A resume is a few pages of text; a 10 MB one is a portfolio, a scan at print resolution,
or the wrong file. Far below `media.MAX_UPLOAD_BYTES`, deliberately — that ceiling is sized for
a reference video and applying it here would let someone put 200 MB of anything in the store.
"""

RESUME_DIRNAME = "resumes"
"""
Resumes live in their own directory under the media root, not beside reference clips.

**Not tidiness.** `media.store_upload` validates against video, image and audio suffixes and
refuses everything else, which is correct for its job — a reference is a face or a voice.
Widening that allowlist so a PDF could pass would mean a document could be uploaded as a face,
and the error message that guards it would start lying. Separate directory, separate writer,
separate size ceiling: the two things are not the same kind of object and the code should not
pretend otherwise.
"""


def store_resume(data: bytes, filename: str) -> Path:
    """
    Write a resume under `<media root>/resumes/` and return the path.

    The uploader's filename supplies only the suffix. Everything else about it -- directories,
    `..`, a leading dot -- is a path a stranger chose on our disk, which is the same reason
    `media.store_upload` discards it and the store refuses ids containing separators.
    """
    import uuid

    from avatar import media

    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED:
        raise ResumeUnreadable(
            f"{suffix or 'that file'} is not a resume format this accepts. "
            f"Supported: {', '.join(SUPPORTED)}."
        )
    if not data:
        raise ResumeUnreadable("the upload was empty.")
    if len(data) > MAX_RESUME_BYTES:
        raise ResumeUnreadable(
            f"the file is {len(data) / 1_048_576:.1f} MB, over the "
            f"{MAX_RESUME_BYTES // 1_048_576} MB ceiling for a resume."
        )

    directory = media.MEDIA_ROOT / RESUME_DIRNAME
    directory.mkdir(parents=True, exist_ok=True)
    target = directory / f"cv-{uuid.uuid4().hex[:12]}{suffix}"
    target.write_bytes(data)
    return target


class ResumeUnreadable(ValueError):
    """
    The file could not be turned into text, with a reason an operator can act on.

    Its own type because the three causes need different answers: an unsupported format means
    convert it, a parser that is not installed means install the extra, and a PDF of scanned
    images means it needs OCR or a text version. Collapsing them into "could not read resume"
    produces a support conversation.
    """


@dataclass(frozen=True)
class Extracted:
    text: str
    chars: int
    """Characters kept. Compare with `total_chars` to see whether it was truncated."""
    total_chars: int
    pages: int | None
    kind: str

    @property
    def truncated(self) -> bool:
        return self.total_chars > self.chars


def _tidy(text: str) -> str:
    """
    Collapse the whitespace a PDF extractor leaves behind, without joining paragraphs.

    PDF text extraction emits a newline per laid-out line, so a wrapped bullet arrives as three
    lines and a two-column layout interleaves them. Collapsing single newlines into spaces while
    keeping blank lines as paragraph breaks is the cheap fix that makes the result readable; it
    does not fix two-column layouts, and nothing here pretends to.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # A single newline between two non-empty lines becomes a space; a blank line stays.
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    return re.sub(r" {2,}", " ", text).strip()


def _from_pdf(data: bytes) -> tuple[str, int | None]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        raise ResumeUnreadable(
            "this is a PDF and the PDF parser is not installed. Install the resume extra: "
            "pip install -e '.[resume]' — or upload a .txt, .md or .docx version."
        ) from None
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ResumeUnreadable(
            f"the PDF could not be parsed ({type(exc).__name__}). It may be corrupt or "
            "password-protected."
        ) from None
    text = "\n\n".join(pages)
    if not text.strip():
        raise ResumeUnreadable(
            f"the PDF has {len(pages)} page(s) but no extractable text, which usually means "
            "it is a scan. It needs OCR, or a text version of the same document."
        )
    return text, len(pages)


def _from_docx(data: bytes) -> tuple[str, int | None]:
    try:
        import docx
    except ModuleNotFoundError:
        raise ResumeUnreadable(
            "this is a Word document and the parser is not installed. Install the resume "
            "extra: pip install -e '.[resume]' — or upload a .txt, .md or .pdf version."
        ) from None
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeUnreadable(
            f"the Word document could not be parsed ({type(exc).__name__}). If it is a .doc "
            "rather than a .docx, save it as .docx first."
        ) from None
    # Tables carry real content on plenty of resumes -- skills grids, dated role lists -- and a
    # paragraph-only read silently drops them, which looks like a thin resume rather than a
    # partial parse.
    blocks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" · ".join(cells))
    text = "\n".join(block for block in blocks if block.strip())
    if not text.strip():
        raise ResumeUnreadable("the Word document contains no text.")
    return text, None


def _from_text(data: bytes) -> tuple[str, int | None]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
        if text.strip():
            return text, None
        raise ResumeUnreadable("the file is empty.")
    raise ResumeUnreadable("the file is not readable as text in UTF-8, UTF-16 or Latin-1.")


def extract(data: bytes, filename: str) -> Extracted:
    """
    Text from a resume, by file extension. Raises `ResumeUnreadable` with a reason.

    Dispatch on the extension rather than on sniffed content, deliberately: the operator chose
    the file and the extension is their statement of what it is. Sniffing would let a
    mislabelled file succeed in a way that is harder to explain than a clear "convert this
    first".
    """
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        raw, pages = _from_pdf(data)
    elif suffix == ".docx":
        raw, pages = _from_docx(data)
    elif suffix in (".txt", ".md", ".markdown"):
        raw, pages = _from_text(data)
    elif suffix == ".doc":
        raise ResumeUnreadable(
            "the old .doc format is not supported. Save it as .docx or export a PDF."
        )
    else:
        raise ResumeUnreadable(
            f"{suffix or 'that file'} is not a resume format this reads. "
            f"Supported: {', '.join(SUPPORTED)}."
        )

    text = _tidy(raw)
    total = len(text)
    if total > MAX_BRIEFING_CHARS:
        # Cut on a paragraph boundary if one is near the limit, so the briefing does not end
        # mid-sentence -- a model handed a truncated sentence tends to complete it as an
        # assumption.
        window = text[:MAX_BRIEFING_CHARS]
        boundary = window.rfind("\n\n")
        text = window[:boundary] if boundary > MAX_BRIEFING_CHARS * 0.7 else window
    return Extracted(
        text=text,
        chars=len(text),
        total_chars=total,
        pages=pages,
        kind=suffix.lstrip("."),
    )


def briefing(candidate: dict[str, object]) -> str:
    """
    The block appended to an interviewer's system prompt for this candidate.

    **The framing is the important part, not the text.** A model handed a resume treats it as
    ground truth and will state its contents back as fact — "you led the migration at X" — which
    is exactly the failure an interview exists to avoid. So the resume arrives explicitly
    labelled as the candidate's own unverified claims, with an instruction to probe rather than
    accept.
    Without that framing the resume makes the interview worse, not better.

    Returns an empty string when there is nothing to say, so the caller can concatenate
    unconditionally.
    """
    name = str(candidate.get("name") or "").strip()
    role = str(candidate.get("role") or "").strip()
    # Two shapes reach here and both are legitimate. The Postgres reader returns `doc` merged
    # into the top level, so `resume_text` is a sibling of `name`; a raw row read straight from
    # the store has it nested under `doc`. Checking both is cheaper than making one backend's
    # reader lie about its own shape.
    nested = candidate.get("doc")
    fields: dict[str, object] = dict(nested) if isinstance(nested, dict) else {}
    text = str(candidate.get("resume_text") or fields.get("resume_text") or "").strip()
    notes = str(candidate.get("notes") or fields.get("notes") or "").strip()

    if not (name or role or text or notes):
        return ""

    parts = ["\n\n--- About the person you are interviewing ---"]
    if name:
        parts.append(f"Their name is {name}. Use it once, early, and then stop using it.")
    if role:
        parts.append(f"They are interviewing for: {role}.")
    if notes:
        parts.append(f"Note from the hiring team: {notes}")
    if text:
        parts.append(
            "Below is the resume they submitted. Treat every line of it as a claim they have "
            "made about themselves, not as established fact. Your job is to find out which "
            "parts are real and how deep they go — use it to choose what to probe, and never "
            "state its contents back to them as though you know it true. If something in it is "
            "central to the role, ask about that specifically rather than asking them to walk "
            "through their background.\n\n"
            f"RESUME AS SUBMITTED:\n{text}"
        )
    return "\n\n".join(parts)
